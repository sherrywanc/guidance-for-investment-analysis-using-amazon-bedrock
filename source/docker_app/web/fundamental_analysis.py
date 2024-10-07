import streamlit as st
from genAI.agents.financial_analysis import analyze_financials
from docx import Document
import io

def render():
    st.subheader("Fundamental Analysis using LLM")
    st.write("Enter a stock ticker below to start the analysis.")

    ticker = st.text_input("Stock Ticker (e.g., AMZN, ZM):", "")
    analyze_button = st.button("Analyze Financials", key="analyze_btn")

    if analyze_button and ticker:
        with st.spinner(f"Analyzing financials for {ticker}..."):
            try:
                user_input = f"Can you analyze the yearly financials for {ticker}?"
                final_output, parsed_data = analyze_financials(user_input)

                st.session_state['final_output'] = final_output
                st.session_state['parsed_data'] = parsed_data
                st.session_state['ticker'] = ticker

                st.success("Analysis complete. Check the results below.")
                st.markdown(f"**Stock Ticker:** {ticker}")
                st.markdown(final_output)

            except Exception as e:
                st.error(f"An error occurred during analysis: {e}")
                st.write("Please check your input or try again later.")
    elif analyze_button:
        st.warning("Please enter a valid stock ticker.")
    st.markdown("---")  # Horizontal line for separation

    # Provide option to save analysis as Word or Text file
    if 'final_output' in st.session_state and 'ticker' in st.session_state:
        if st.button("Save Analysis as Word"):
            doc = Document()
            doc.add_heading(f"Financial Analysis for {st.session_state['ticker']}", 0)
            doc.add_paragraph(st.session_state['final_output'])
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            st.download_button(
                label="Download Word Document",
                data=doc_io,
                file_name=f"{st.session_state['ticker']}_financial_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        if st.button("Save Analysis as Text"):
            text_io = io.StringIO(st.session_state['final_output'])
            st.download_button(
                label="Download Text File",
                data=text_io,
                file_name=f"{st.session_state['ticker']}_financial_analysis.txt",
                mime="text/plain"
            )
